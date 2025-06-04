from flask import Flask, request, jsonify, abort
from flask_cors import CORS
import torch
from transformers import (
    RobertaTokenizerFast, RobertaForTokenClassification,
    DetrForObjectDetection, DetrImageProcessor, DetrConfig,
    PegasusTokenizer, PegasusForConditionalGeneration
)
from safetensors.torch import safe_open
from PIL import Image, ImageDraw, ImageFont
import io
import base64
import os
import numpy as np
import onnxruntime as ort
from torchvision import transforms
from docx import Document
import pdfplumber
import pytesseract
import traceback
import re
import logging
import tempfile

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Setup logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Device setup
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
logger.info(f"Using device: {device}")

# Model paths (UPDATE THESE PATHS)
SERVER_DIR = os.path.dirname(__file__)
CLASSIFICATION_MODEL_PATH = r"C:\All Files\gpt\server\classification\vit_model.onnx"  # Confirm correct path
DETR_MODEL_DIR = r"C:\All Files\gpt\server\detr_model"  # Confirm correct path
PII_MODEL_PATH = r"C:\All Files\gpt\server\fine_tuned_roberta_pii"  # Confirm correct path
SUMMARIZATION_MODEL_PATH = r"C:\Users\jashw\Downloads\pegasus_finetuned_2\pegasus_finetuned"  # Confirm correct path
TESSERACT_PATH = os.environ.get("TESSERACT_PATH", r"C:\Program Files\Tesseract-OCR\tesseract.exe")

# Set Tesseract path
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

# Load models
def load_models():
    global classification_ort_session, classification_transform
    global detr_model, detr_processor
    global pii_tokenizer, pii_model
    global pegasus_tokenizer, pegasus_model

    # Document Classification
    if not os.path.exists(CLASSIFICATION_MODEL_PATH):
        raise FileNotFoundError(f"ONNX model not found at: {CLASSIFICATION_MODEL_PATH}")
    classification_ort_session = ort.InferenceSession(CLASSIFICATION_MODEL_PATH)
    classification_transform = transforms.Compose([
        transforms.Resize((224, 224)),
        transforms.ToTensor(),
        transforms.Normalize(mean=[0.5, 0.5, 0.5], std=[0.5, 0.5, 0.5])
    ])
    logger.info("Classification model loaded")

    # Image Masking
    config_path = os.path.join(DETR_MODEL_DIR, "config.json")
    safetensors_path = os.path.join(DETR_MODEL_DIR, "model.safetensors")
    if not os.path.exists(config_path) or not os.path.exists(safetensors_path):
        raise FileNotFoundError(f"DETR model files missing at: {DETR_MODEL_DIR}")
    config = DetrConfig.from_json_file(config_path)
    detr_model = DetrForObjectDetection(config)
    with safe_open(safetensors_path, framework="pt", device="cpu") as f:
        state_dict = {k: f.get_tensor(k) for k in f.keys()}
    detr_model.load_state_dict(state_dict)
    detr_model.eval()
    detr_processor = DetrImageProcessor.from_pretrained("facebook/detr-resnet-50")
    logger.info("Image Masking model loaded")

    # Text PII Masking
    if not os.path.exists(PII_MODEL_PATH):
        raise FileNotFoundError(f"PII model not found at: {PII_MODEL_PATH}")
    pii_tokenizer = RobertaTokenizerFast.from_pretrained(PII_MODEL_PATH)
    pii_model = RobertaForTokenClassification.from_pretrained(PII_MODEL_PATH)
    pii_model = pii_model.to(device)
    logger.info("PII Masking model loaded")

    # Document Summarization
    if not os.path.exists(SUMMARIZATION_MODEL_PATH):
        raise FileNotFoundError(f"Summarization model not found at: {SUMMARIZATION_MODEL_PATH}")
    pegasus_tokenizer = PegasusTokenizer.from_pretrained(SUMMARIZATION_MODEL_PATH)
    pegasus_model = PegasusForConditionalGeneration.from_pretrained(SUMMARIZATION_MODEL_PATH).to(device)
    logger.info("Summarization model loaded")

# Load models at startup
try:
    load_models()
except Exception as e:
    logger.error(f"Error loading models: {e}")
    raise

# Classification variables
SELECTED_CLASSES = ["budget", "form", "invoice"]

# PII Masking variables
labels_list = [
    'O', 'B-AADHAAR_ID', 'B-ACCOUNTNAME', 'B-ACCOUNTNUMBER', 'B-ADDRESS', 'B-AGE', 'B-AMOUNT', 'B-BANK', 'B-BBAN', 'B-BIC',
    'B-BITCOINADDRESS', 'B-BUILDINGNUMBER', 'B-CITY', 'B-COMPANY_NAME', 'B-COUNTY', 'B-CREDITCARDCVV', 'B-CREDITCARDISSUER',
    'B-CREDITCARDNUMBER', 'B-CURRENCY', 'B-CURRENCYCODE', 'B-CURRENCYNAME', 'B-CURRENCYSYMBOL', 'B-DATE', 'B-DATE_OF_BIRTH',
    'B-DRIVER_LICENSE', 'B-EMAIL', 'B-ETHEREUMADDRESS', 'B-FirstNAME', 'B-FULLNAME', 'B-GENDER', 'B-IBAN', 'B-IP', 'B-IPV4',
    'B-IPV6', 'B-JOBAREA', 'B-JOBDESCRIPTOR', 'B-JOBTITLE', 'B-JOBTYPE', 'B-LASTNAME', 'B-LATITUDE', 'B-LICENSE_PLATE',
    'B-LITECOINADDRESS', 'B-LONGITUDE', 'B-MAC', 'B-MASKEDNUMBER', 'B-MIDDLENAME', 'B-PAN_NUMBER', 'B-PASSWORD', 'B-PHONEIMEI',
    'B-PHONE_NUMBER', 'B-PIN', 'B-PREFIX', 'B-SECONDARYADDRESS', 'B-SEX', 'B-SSN', 'B-STATE', 'B-STREET', 'B-STREETADDRESS',
    'B-SUFFIX', 'B-TIME', 'B-URL', 'B-USERAGENT', 'B-USERNAME', 'B-VEHICLEVIN', 'B-VEHICLEVRM', 'B-ZIPCODE', 'I-AADHAAR_ID',
    'I-ACCOUNTNAME', 'I-ACCOUNTNUMBER', 'I-ADDRESS', 'I-AGE', 'I-AMOUNT', 'I-BANK', 'I-BBAN', 'I-BIC', 'I-BITCOINADDRESS',
    'I-BUILDINGNUMBER', 'I-CITY', 'I-COMPANY_NAME', 'I-CREDITCARDCVV', 'I-CREDITCARDISSUER', 'I-CREDITCARDNUMBER', 'I-CURRENCY',
    'I-CURRENCYCODE', 'I-CURRENCYNAME', 'I-CURRENCYSYMBOL', 'I-DATE', 'I-DATE_OF_BIRTH', 'I-DRIVER_LICENSE', 'I-EMAIL',
    'I-ETHEREUMADDRESS', 'I-FirstNAME', 'I-FULLNAME', 'I-GENDER', 'I-IBAN', 'I-IP', 'I-IPV4', 'I-IPV6', 'I-JOBAREA', 'I-JOBTITLE',
    'I-JOBTYPE', 'I-LASTNAME', 'I-LATITUDE', 'I-LICENSE_PLATE', 'I-LITECOINADDRESS', 'I-LONGITUDE', 'I-MAC', 'I-MASKEDNUMBER',
    'I-MIDDLENAME', 'I-PAN_NUMBER', 'I-PASSWORD', 'I-PHONEIMEI', 'I-PHONE_NUMBER', 'I-PIN', 'I-PREFIX', 'I-SECONDARYADDRESS',
    'I-SSN', 'I-STATE', 'I-STREET', 'I-STREETADDRESS', 'I-SUFFIX', 'I-TIME', 'I-URL', 'I-USERAGENT', 'I-USERNAME', 'I-VEHICLEVIN',
    'I-VEHICLEVRM', 'I-ZIPCODE', 'O'
]
ENTITY_MASKS = {
    "AADHAAR_ID": "{Aadhaar_ID}", "ACCOUNTNAME": "{Account_Name}", "ACCOUNTNUMBER": "{Account_Number}",
    "ADDRESS": "{Address}", "AGE": "{Age}", "AMOUNT": "{Amount}", "BANK": "{Bank}", "BBAN": "{BBAN}", "BIC": "{BIC}",
    "BITCOINADDRESS": "{Bitcoin_Address}", "BUILDINGNUMBER": "{Building_Number}", "CITY": "{City}",
    "COMPANY_NAME": "{Company_Name}", "CREDITCARDCVV": "{Card_CVV}", "CREDITCARDNUMBER": "{Card_Number}",
    "CURRENCY": "{Currency}", "DATE": "{Date}", "DATE_OF_BIRTH": "{DOB}", "DRIVER_LICENSE": "{Driver_License}",
    "EMAIL": "{E-Mail}", "ETHEREUMADDRESS": "{Ethereum_Address}", "FIRSTNAME": "{First_Name}", "FULLNAME": "{Full_Name}",
    "GENDER": "{Gender}", "IBAN": "{IBAN}", "IP": "{IP}", "PHONE_NUMBER": "{Phone_Number}", "PAN_NUMBER": "{PAN_Number}",
    "PASSWORD": "{Password}", "SSN": "{SSN}", "STATE": "{State}", "STREETADDRESS": "{Street_Address}", "URL": "{URL}",
    "USERNAME": "{Username}", "VEHICLEVIN": "{Vehicle_VIN}", "ZIPCODE": "{ZIP_Code}"
}

# Helper functions
def render_text_to_image(text, width=800, height=1000):
    image = Image.new("RGB", (width, height), color="white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 18)
    except:
        font = ImageFont.load_default()
    lines = text.split("\n")
    y = 10
    for line in lines:
        if y > height - 30:
            break
        draw.text((10, y), line[:120], fill="black", font=font)
        y += 25
    return image

def extract_text_from_docx(file_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp.close()
        doc = Document(tmp.name)
        os.unlink(tmp.name)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_from_pdf(file_bytes):
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        return text

def mask_predictions(image, boxes, fill_color=(0, 0, 0)):
    masked_image = image.copy()
    draw = ImageDraw.Draw(masked_image)
    for box in boxes:
        xmin, ymin, xmax, ymax = box.tolist()
        draw.rectangle([xmin, ymin, xmax, ymax], fill=fill_color)
    return masked_image

def predict_ner(model, tokenizer, example):
    logger.info("Input text: %s", example)
    encoding = tokenizer(example, return_offsets_mapping=True, truncation=True, padding="max_length", max_length=512, return_attention_mask=True)
    tokens_ids = encoding['input_ids']
    attention_mask = encoding['attention_mask']
    offset_mapping = encoding['offset_mapping']
    tokens = tokenizer.convert_ids_to_tokens(tokens_ids)
    input_ids_tensor = torch.tensor(tokens_ids).unsqueeze(0).to(device)
    attention_mask_tensor = torch.tensor(attention_mask).unsqueeze(0).to(device)
    with torch.no_grad():
        outputs = model(input_ids_tensor, attention_mask=attention_mask_tensor)
    preds = torch.argmax(outputs.logits, dim=2).cpu().numpy().flatten()
    entity_spans = {}
    current_entity = ""
    entity_type = None
    entity_start = None
    for i, (token, pred) in enumerate(zip(tokens, preds)):
        label = labels_list[pred]
        start, end = offset_mapping[i]
        if start == end:
            continue
        if label.startswith("B-"):
            if current_entity:
                entity_spans.setdefault(entity_type, []).append((current_entity, entity_start))
            current_entity = example[start:end]
            entity_type = label[2:]
            entity_start = start
        elif label.startswith("I-") and entity_type == label[2:]:
            current_entity += example[start:end]
        else:
            if current_entity:
                entity_spans.setdefault(entity_type, []).append((current_entity, entity_start))
                current_entity = ""
                entity_type = None
    if current_entity:
        entity_spans.setdefault(entity_type, []).append((current_entity, entity_start))
    logger.info("Detected entities: %s", entity_spans)
    masked_text = example
    for entity_type, spans in entity_spans.items():
        mask = ENTITY_MASKS.get(entity_type, "{MASKED}")
        for entity, _ in sorted(spans, key=lambda x: x[1], reverse=True):
            masked_text = re.sub(re.escape(entity), mask, masked_text, count=1)
    logger.info("Masked text: %s", masked_text)
    return masked_text

def extract_text_from_file(file, content: bytes) -> str:
    filename = file.filename.lower()
    if filename.endswith(".txt"):
        return content.decode("utf-8")
    elif filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])
    elif filename.endswith(".pdf"):
        return extract_text_from_pdf(content)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        image = Image.open(io.BytesIO(content)).convert("RGB")
        return pytesseract.image_to_string(image)
    else:
        raise ValueError("Unsupported file type. Allowed: pdf, docx, txt, jpg, jpeg, png")

# Routes
@app.route("/")
def home():
    return jsonify({"message": "FinGPT Unified Server"})

@app.route("/health")
def health_check():
    return jsonify({
        "status": "ok",
        "tasks": ["classification", "image-masking", "pii", "summarization"]
    })

# Document Classification Routes
@app.route("/classification/predict", methods=["POST"])
def classification_predict():
    logger.info(f"Received request to /classification/predict from: {request.remote_addr}")
    if "file" not in request.files:
        logger.error("No file provided")
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    filename = file.filename.lower()
    extension = filename.rsplit(".", 1)[-1] if "." in filename else ""
    logger.info(f"Processing file: {filename}, extension: {extension}")
    if extension not in ["jpg", "jpeg", "png", "bmp", "txt", "docx", "pdf"]:
        logger.error(f"Unsupported file type: {extension}")
        return jsonify({"error": "Unsupported file type"}), 400
    try:
        contents = file.read()
        if extension in ["jpg", "jpeg", "png", "bmp"]:
            logger.info("Processing image file")
            image = Image.open(io.BytesIO(contents)).convert("RGB")
        elif extension == "txt":
            logger.info("Processing text file")
            text = contents.decode("utf-8")
            image = render_text_to_image(text)
        elif extension == "docx":
            logger.info("Processing DOCX file")
            text = extract_text_from_docx(contents)
            image = render_text_to_image(text)
        elif extension == "pdf":
            logger.info("Processing PDF file")
            text = extract_text_from_pdf(contents)
            image = render_text_to_image(text)
        logger.info("Running model inference")
        image_tensor = classification_transform(image).unsqueeze(0).numpy()
        input_name = classification_ort_session.get_inputs()[0].name
        ort_inputs = {input_name: image_tensor}
        ort_outs = classification_ort_session.run(None, ort_inputs)
        logits = ort_outs[0]
        probabilities = torch.nn.functional.softmax(torch.from_numpy(logits), dim=1).numpy().flatten()
        predicted_class = int(np.argmax(probabilities))
        confidence = float(probabilities[predicted_class])
        logger.info(f"Prediction: {SELECTED_CLASSES[predicted_class]}, Confidence: {confidence:.2f}")
        return jsonify({
            "prediction": SELECTED_CLASSES[predicted_class],
            "confidence": round(confidence, 2)
        })
    except Exception as e:
        logger.error(f"Classification error: {e}")
        return jsonify({"error": str(e)}), 500

# Image Masking Routes
@app.route("/predict", methods=["POST"])
def image_masking_predict():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    try:
        image = Image.open(file.stream).convert("RGB")
        inputs = detr_processor(images=image, return_tensors="pt")
        with torch.no_grad():
            outputs = detr_model(**inputs)
        target_sizes = torch.tensor([image.size[::-1]])
        results = detr_processor.post_process_object_detection(outputs, target_sizes=target_sizes, threshold=0.9)[0]
        masked_image = mask_predictions(image, results["boxes"])
        buffered = io.BytesIO()
        masked_image.save(buffered, format="PNG")
        masked_image_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
        response = {
            "boxes": results["boxes"].tolist(),
            "labels": [detr_model.config.id2label.get(label.item(), f"Class {label.item()}") for label in results["labels"]],
            "scores": results["scores"].tolist(),
            "masked_image": f"data:image/png;base64,{masked_image_base64}"
        }
        return jsonify(response)
    except Exception as e:
        logger.error(f"Image Masking error: {e}")
        return jsonify({"error": str(e)}), 500

# Text PII Masking Routes
@app.route("/mask", methods=["POST"])
def pii_mask():
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({'error': 'No text provided'}), 400
        text = data['text']
        if not text.strip():
            return jsonify({'error': 'Text is empty'}), 400
        masked_text = predict_ner(pii_model, pii_tokenizer, text)
        return jsonify({'masked_text': masked_text})
    except Exception as e:
        logger.error(f"PII Masking error: {e}")
        return jsonify({'error': str(e)}), 500

# Document Summarization Routes
@app.route("/summarize", methods=["POST"])
def summarize():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400
        file = request.files['file']
        allowed_extensions = (".pdf", ".docx", ".txt", ".jpeg", ".jpg", ".png")
        if not file.filename.lower().endswith(allowed_extensions):
            return jsonify({"error": "Only PDF, DOCX, TXT, JPG, JPEG, PNG files are allowed."}), 400
        content = file.read()
        text = extract_text_from_file(file, content)
        if not text or not text.strip():
            return jsonify({"error": "No readable text found in the uploaded file."}), 400
        encoded = pegasus_tokenizer(
            text,
            return_tensors="pt",
            truncation=True,
            padding="longest"
        ).to(device)
        with torch.no_grad():
            summary_ids = pegasus_model.generate(
                input_ids=encoded["input_ids"],
                attention_mask=encoded["attention_mask"],
                max_length=100,
                num_beams=5,
                early_stopping=True
            )
        summary = pegasus_tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        return jsonify({"summary": summary})
    except Exception as e:
        logger.error(f"Summarization error: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    logger.info("Starting FinGPT Unified Server")
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))